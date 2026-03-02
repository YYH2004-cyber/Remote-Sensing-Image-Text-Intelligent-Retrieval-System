import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def markdown_to_word(md_file_path, word_file_path):
    doc = Document()
    
    with open(md_file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        elif line.startswith('##### '):
            doc.add_heading(line[6:], level=5)
        elif line.startswith('###### '):
            doc.add_heading(line[7:], level=6)
        elif line.strip().startswith('- '):
            p = doc.add_paragraph(line.strip()[2:], style='List Bullet')
        elif line.strip().startswith('* '):
            p = doc.add_paragraph(line.strip()[2:], style='List Bullet')
        elif line.strip().startswith('1. '):
            p = doc.add_paragraph(line.strip()[3:], style='List Number')
        elif line.strip().startswith('2. '):
            p = doc.add_paragraph(line.strip()[3:], style='List Number')
        elif line.strip().startswith('3. '):
            p = doc.add_paragraph(line.strip()[3:], style='List Number')
        elif line.strip().startswith('4. '):
            p = doc.add_paragraph(line.strip()[3:], style='List Number')
        elif line.strip().startswith('5. '):
            p = doc.add_paragraph(line.strip()[3:], style='List Number')
        elif line.strip().startswith('6. '):
            p = doc.add_paragraph(line.strip()[3:], style='List Number')
        elif line.strip().startswith('7. '):
            p = doc.add_paragraph(line.strip()[3:], style='List Number')
        elif line.strip().startswith('8. '):
            p = doc.add_paragraph(line.strip()[3:], style='List Number')
        elif line.strip().startswith('9. '):
            p = doc.add_paragraph(line.strip()[3:], style='List Number')
        elif line.strip().startswith('0. '):
            p = doc.add_paragraph(line.strip()[3:], style='List Number')
        elif line.strip().startswith('Q:'):
            p = doc.add_paragraph(line.strip(), style='List Bullet')
            run = p.runs[0]
            run.bold = True
            run.font.color.rgb = RGBColor(0, 51, 102)
        elif line.strip().startswith('A:'):
            p = doc.add_paragraph(line.strip(), style='List Bullet')
            run = p.runs[0]
            run.bold = True
            run.font.color.rgb = RGBColor(0, 102, 51)
        elif line.strip().startswith('✅'):
            p = doc.add_paragraph(line.strip(), style='List Bullet')
            run = p.runs[0]
            run.font.color.rgb = RGBColor(0, 128, 0)
        elif line.strip().startswith('❌'):
            p = doc.add_paragraph(line.strip(), style='List Bullet')
            run = p.runs[0]
            run.font.color.rgb = RGBColor(255, 0, 0)
        elif line.strip().startswith('⚠️'):
            p = doc.add_paragraph(line.strip(), style='List Bullet')
            run = p.runs[0]
            run.font.color.rgb = RGBColor(255, 165, 0)
        elif line.strip().startswith('💡'):
            p = doc.add_paragraph(line.strip(), style='List Bullet')
            run = p.runs[0]
            run.font.color.rgb = RGBColor(0, 102, 204)
        elif line.strip().startswith('📚'):
            p = doc.add_paragraph(line.strip(), style='List Bullet')
            run = p.runs[0]
            run.font.color.rgb = RGBColor(128, 0, 128)
        elif line.strip().startswith('🎯'):
            p = doc.add_paragraph(line.strip(), style='List Bullet')
            run = p.runs[0]
            run.font.color.rgb = RGBColor(204, 0, 0)
        elif line.strip().startswith('📝'):
            p = doc.add_paragraph(line.strip(), style='List Bullet')
            run = p.runs[0]
            run.font.color.rgb = RGBColor(0, 0, 204)
        elif line.strip().startswith('🔧'):
            p = doc.add_paragraph(line.strip(), style='List Bullet')
            run = p.runs[0]
            run.font.color.rgb = RGBColor(128, 128, 0)
        elif line.strip().startswith('🚀'):
            p = doc.add_paragraph(line.strip(), style='List Bullet')
            run = p.runs[0]
            run.font.color.rgb = RGBColor(255, 0, 255)
        elif line.strip().startswith('⭐'):
            p = doc.add_paragraph(line.strip(), style='List Bullet')
            run = p.runs[0]
            run.font.color.rgb = RGBColor(255, 215, 0)
        elif line.strip().startswith('📊'):
            p = doc.add_paragraph(line.strip(), style='List Bullet')
            run = p.runs[0]
            run.font.color.rgb = RGBColor(0, 128, 128)
        elif line.strip().startswith('🔍'):
            p = doc.add_paragraph(line.strip(), style='List Bullet')
            run = p.runs[0]
            run.font.color.rgb = RGBColor(0, 0, 128)
        elif line.strip().startswith('💻'):
            p = doc.add_paragraph(line.strip(), style='List Bullet')
            run = p.runs[0]
            run.font.color.rgb = RGBColor(128, 0, 64)
        elif line.strip().startswith('📖'):
            p = doc.add_paragraph(line.strip(), style='List Bullet')
            run = p.runs[0]
            run.font.color.rgb = RGBColor(0, 64, 64)
        elif line.strip().startswith('🎓'):
            p = doc.add_paragraph(line.strip(), style='List Bullet')
            run = p.runs[0]
            run.font.color.rgb = RGBColor(128, 64, 0)
        elif line.strip().startswith('📋'):
            p = doc.add_paragraph(line.strip(), style='List Bullet')
            run = p.runs[0]
            run.font.color.rgb = RGBColor(64, 64, 64)
        elif line.strip().startswith('🌟'):
            p = doc.add_paragraph(line.strip(), style='List Bullet')
            run = p.runs[0]
            run.font.color.rgb = RGBColor(255, 192, 203)
        elif line.strip().startswith('📌'):
            p = doc.add_paragraph(line.strip(), style='List Bullet')
            run = p.runs[0]
            run.font.color.rgb = RGBColor(255, 140, 0)
        elif line.strip().startswith('🔄'):
            p = doc.add_paragraph(line.strip(), style='List Bullet')
            run = p.runs[0]
            run.font.color.rgb = RGBColor(0, 128, 128)
        elif line.strip().startswith('🎨'):
            p = doc.add_paragraph(line.strip(), style='List Bullet')
            run = p.runs[0]
            run.font.color.rgb = RGBColor(147, 112, 219)
        elif line.strip().startswith('🔬'):
            p = doc.add_paragraph(line.strip(), style='List Bullet')
            run = p.runs[0]
            run.font.color.rgb = RGBColor(75, 0, 130)
        elif line.strip().startswith('📈'):
            p = doc.add_paragraph(line.strip(), style='List Bullet')
            run = p.runs[0]
            run.font.color.rgb = RGBColor(34, 139, 34)
        elif line.strip().startswith('🔗'):
            p = doc.add_paragraph(line.strip(), style='List Bullet')
            run = p.runs[0]
            run.font.color.rgb = RGBColor(0, 0, 255)
        elif line.strip().startswith('---'):
            doc.add_paragraph().add_run().add_break()
        elif line.strip() == '':
            doc.add_paragraph()
        elif line.strip().startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph()
            run = p.add_run(code_text)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            p.paragraph_format.left_indent = Inches(0.5)
        elif line.strip().startswith('>'):
            p = doc.add_paragraph(line.strip()[1:])
            run = p.runs[0]
            run.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128)
        elif line.strip().startswith('**'):
            text = line.strip().strip('*')
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
        elif line.strip().startswith('*'):
            text = line.strip().strip('*')
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.italic = True
        elif line.strip().startswith('['):
            match = re.match(r'\[(.*?)\]\((.*?)\)', line.strip())
            if match:
                link_text = match.group(1)
                link_url = match.group(2)
                p = doc.add_paragraph()
                run = p.add_run(link_text)
                run.font.color.rgb = RGBColor(0, 0, 255)
                run.underline = True
        else:
            if line.strip():
                doc.add_paragraph(line.strip())
        
        i += 1
    
    doc.save(word_file_path)
    print(f"Word 文档已生成: {word_file_path}")

if __name__ == "__main__":
    markdown_to_word(
        "/home/yyh2004/demo/GRADUATION_AND_INTERVIEW_PREP.md",
        "/home/yyh2004/demo/GRADUATION_AND_INTERVIEW_PREP.docx"
    )
